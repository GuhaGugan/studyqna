from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any, Optional
import io
import html
import unicodedata
import tempfile

# Optional import for pdfkit (requires wkhtmltopdf binary)
try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False
    pdfkit = None

# Optional import for xhtml2pdf (pure Python, good rendering)
try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    XHTML2PDF_AVAILABLE = False
    pisa = None

# Optional import for Playwright (best rendering quality - uses Chromium)
try:
    from playwright.async_api import async_playwright
    PLAYWRIGHT_AVAILABLE = True
    print("✅ Playwright is available - will use for PDF generation")
except ImportError as e:
    PLAYWRIGHT_AVAILABLE = False
    async_playwright = None
    print(f"⚠️  Playwright not available: {e}. Install with: pip install playwright && playwright install chromium")
from app.font_manager import (
    get_font_for_language,
    detect_language,
    is_rtl,
    get_alignment,
    initialize_fonts
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

# Initialize fonts on module import
try:
    initialize_fonts()
except Exception as e:
    print(f"⚠️  Font initialization warning: {e}")

# Font name mapping: short names (for use in <font> tags) -> file names
FONT_NAME_MAP = {
    "NotoLatin": "NotoSans-Regular.ttf",
    "NotoTamil": "NotoSansTamil-Regular.ttf",
    "NotoDeva": "NotoSansDevanagari-Regular.ttf",
    "NotoTelugu": "NotoSansTelugu-Regular.ttf",
    "NotoKannada": "NotoSansKannada-Regular.ttf",
    "NotoMalayalam": "NotoSansMalayalam-Regular.ttf",
    "NotoArabic": "NotoSansArabic-Regular.ttf",
}

# Language to short font name mapping (for use in <font name='...'> tags)
LANGUAGE_TO_FONT_NAME = {
    "english": "NotoLatin",
    "spanish": "NotoLatin",
    "tamil": "NotoTamil",
    "hindi": "NotoDeva",
    "telugu": "NotoTelugu",
    "kannada": "NotoKannada",
    "malayalam": "NotoMalayalam",
    "arabic": "NotoArabic",
}


def format_structured_answer(answer, marks=0):
    """Convert structured answer (dict) to formatted text string (shared for PDF/DOCX/TXT)."""
    if isinstance(answer, dict):
        # Check for English literature format (introduction, explanation, analysis, conclusion)
        if answer.get("introduction") or answer.get("explanation") or answer.get("analysis") or answer.get("conclusion"):
            parts = []
            if answer.get("introduction"):
                parts.append(f"Introduction: {answer.get('introduction')}")
            if answer.get("explanation"):
                parts.append(f"Explanation: {answer.get('explanation')}")
            if answer.get("analysis"):
                parts.append(f"Analysis: {answer.get('analysis')}")
            if answer.get("conclusion"):
                parts.append(f"Conclusion: {answer.get('conclusion')}")
            return "\n".join(parts)
        
        # Check for Science format (definition, explanation, example, conclusion)
        if answer.get("definition") and (answer.get("explanation") or answer.get("example") or answer.get("conclusion")):
            # Only if it doesn't have math-style fields (to distinguish from math)
            if not (answer.get("formula") or answer.get("given") or answer.get("steps")):
                parts = []
                if answer.get("definition"):
                    parts.append(f"Definition: {answer.get('definition')}")
                if answer.get("explanation"):
                    parts.append(f"Explanation: {answer.get('explanation')}")
                if answer.get("example"):
                    parts.append(f"Example: {answer.get('example')}")
                if answer.get("conclusion"):
                    parts.append(f"Conclusion: {answer.get('conclusion')}")
                return "\n".join(parts)
        
        # Check for Social Science format (background, context, key_points, explanation, conclusion)
        if answer.get("background") or answer.get("context") or answer.get("key_points"):
            parts = []
            if answer.get("background"):
                parts.append(f"Background: {answer.get('background')}")
            if answer.get("context"):
                parts.append(f"Context: {answer.get('context')}")
            if answer.get("key_points"):
                kp = answer.get("key_points")
                if isinstance(kp, list):
                    parts.append(f"Key Points: {' '.join(str(v) for v in kp)}")
                else:
                    parts.append(f"Key Points: {kp}")
            if answer.get("explanation"):
                parts.append(f"Explanation: {answer.get('explanation')}")
            if answer.get("conclusion"):
                parts.append(f"Conclusion: {answer.get('conclusion')}")
            return "\n".join(parts)
        
        # Board-style format (post-processed 10-mark math)
        if (
            answer.get("substitution") is not None
            or answer.get("calculation") is not None
            or answer.get("final_answer") is not None
        ):
            parts = []
            if answer.get("given"):
                parts.append(f"Given: {answer.get('given')}")
            if answer.get("formula"):
                parts.append(f"Formula: {answer.get('formula')}")
            if answer.get("substitution"):
                parts.append(f"Substitution: {answer.get('substitution')}")
            if answer.get("calculation"):
                parts.append(f"Calculation:\n{answer.get('calculation')}")
            if answer.get("final_answer"):
                parts.append(f"{answer.get('final_answer')}")
            return "\n".join(parts)

        # Original structured format (Math and other subjects)
        parts = []
        if answer.get("given"):
            parts.append(f"Given: {answer.get('given')}")
        if answer.get("definition"):
            parts.append(f"Definition: {answer.get('definition')}")
        if answer.get("formula"):
            parts.append(f"Formula/Theorem: {answer.get('formula')}")
        if answer.get("coefficients"):
            parts.append(f"Coefficients: {answer.get('coefficients')}")
        if answer.get("steps") and isinstance(answer.get("steps"), list):
            parts.append("Step-by-step Working:")
            for i, step in enumerate(answer.get("steps", []), 1):
                parts.append(f"Step {i}: {step}")
        if answer.get("function_values") and isinstance(answer.get("function_values"), list):
            parts.append("Function Values:")
            for value in answer.get("function_values", []):
                parts.append(f"  • {value}")
        if answer.get("final"):
            parts.append(f"Final Conclusion: {answer.get('final')}")
        return "\n".join(parts) if parts else ""

    return str(answer) if answer else ""

# Defensive: ensure fonts are manually registered with short names (as per example)
def _ensure_manual_font_registration():
    """Register Noto fonts directly from the fonts directory with short names.
    CRITICAL: TTFont automatically embeds font subset in PDF (prevents black blocks).
    ReportLab's TTFont embeds only used glyphs by default, which is fine for our use case."""
    base_dir = os.path.join(os.path.dirname(__file__), "fonts")
    registered_count = 0
    for short_name, file_name in FONT_NAME_MAP.items():
        path = os.path.join(base_dir, file_name)
        if os.path.exists(path):
            try:
                # TTFont automatically embeds font subset (only used glyphs) in PDF
                # This is efficient and ensures all used characters are available
                font_obj = TTFont(short_name, path)
                pdfmetrics.registerFont(font_obj)
                registered_count += 1
                # Only print on first registration to avoid spam
                if registered_count == 1:
                    print(f"✅ Registered font: {short_name} (and {len(FONT_NAME_MAP)-1} others)")
            except Exception as e:
                print(f"⚠️ Could not register font {short_name}: {e}")
        else:
            print(f"⚠️ Font file not found: {path}")

_ensure_manual_font_registration()

async def _generate_pdf_playwright_async(html_string: str, font_path: str, font_name: str) -> bytes:
    """
    Generate PDF using Playwright (Chromium) - BEST rendering quality.
    Uses real browser engine for pixel-perfect rendering matching web interface.
    Fonts are embedded via base64 in HTML, so no file system access needed.
    Enhanced settings for maximum accuracy and clarity.
    ASYNC VERSION for use with FastAPI.
    """
    print(f"🎭 Playwright: Generating PDF with font: {font_name}")
    try:
        async with async_playwright() as p:
            # Launch browser with high-quality rendering
            browser = await p.chromium.launch(
                headless=True,
                args=[
                    '--disable-web-security',  # Allow font loading
                    '--disable-features=TranslateUI',
                    '--disable-ipc-flooding-protection',
                    '--font-render-hinting=none',  # Better font rendering
                ]
            )
            page = await browser.new_page()
            
            # Set viewport for consistent rendering
            await page.set_viewport_size({"width": 1200, "height": 1600})
            
            # Set content with HTML (fonts are already base64 embedded in HTML)
            print("🎭 Playwright: Setting HTML content...")
            await page.set_content(html_string, wait_until='networkidle')
            
            # Wait for KaTeX to load and render LaTeX
            print("🎭 Playwright: Waiting for KaTeX to render LaTeX...")
            try:
                # Wait for KaTeX to be available
                await page.wait_for_function(
                    "typeof renderMathInElement !== 'undefined' || typeof katex !== 'undefined'",
                    timeout=10000
                )
                # Wait a bit more for KaTeX to finish rendering
                await page.wait_for_timeout(1500)  # KaTeX renders synchronously, but give it time
                print("✅ KaTeX rendering complete")
            except Exception as e:
                print(f"⚠️  KaTeX wait warning: {e} (continuing anyway)")
            
            # Wait for fonts to fully load and render
            print("🎭 Playwright: Waiting for fonts to load...")
            await page.wait_for_timeout(1000)  # Additional wait for font rendering
            
            # Verify font is loaded by checking computed styles
            try:
                font_family = await page.evaluate("getComputedStyle(document.body).fontFamily")
                print(f"🎭 Playwright: Body font-family: {font_family}")
            except Exception as e:
                print(f"⚠️  Could not verify font: {e}")
            
            # Generate PDF with maximum quality settings
            print("🎭 Playwright: Generating PDF...")
            pdf_bytes = await page.pdf(
                format='A4',
                margin={
                    'top': '0.75in',
                    'right': '0.75in',
                    'bottom': '0.75in',
                    'left': '0.75in'
                },
                print_background=True,
                prefer_css_page_size=True,
                scale=1.0,  # 100% scale for maximum clarity
                display_header_footer=False,
            )
            
            await browser.close()
            print(f"✅ Playwright: PDF generated successfully ({len(pdf_bytes)} bytes)")
            return pdf_bytes
    except Exception as e:
        print(f"❌ Playwright error: {e}")
        import traceback
        traceback.print_exc()
        raise

def _generate_pdf_xhtml2pdf(html_string: str, font_path: str, font_name: str) -> bytes:
    """
    Generate PDF using xhtml2pdf (pisa) - pure Python, good rendering quality.
    Better Unicode support than ReportLab, no external dependencies.
    """
    buffer = io.BytesIO()
    
    # Create PDF from HTML string
    result = pisa.CreatePDF(
        html_string,
        dest=buffer,
        encoding='utf-8',
        link_callback=None  # No external links to resolve
    )
    
    if result.err:
        raise Exception(f"xhtml2pdf error: {result.err}")
    
    buffer.seek(0)
    return buffer.read()

def _generate_pdf_reportlab(
    qna_data: Dict[str, Any],
    output_format: str,
    title: str,
    target_language: str,
    normalize_text,
    font_name: str,
    is_rtl_lang: bool,
    text_alignment
) -> bytes:
    """Fallback ReportLab PDF generation when pdfkit is not available.
    CRITICAL: Ensures fonts are properly embedded to prevent black blocks."""
    
    def format_structured_answer(answer, marks=0):
        """Convert structured answer (dict) to formatted text string"""
        if isinstance(answer, dict):
            # Check if it's board-style format (post-processed 10-mark math)
            if answer.get("substitution") is not None or answer.get("calculation") is not None or answer.get("final_answer") is not None:
                # Board-style format: Given, Formula, Substitution, Calculation, Final Answer
                parts = []
                if answer.get("given"):
                    parts.append(f"Given: {answer.get('given')}")
                if answer.get("formula"):
                    parts.append(f"Formula: {answer.get('formula')}")
                if answer.get("substitution"):
                    parts.append(f"Substitution: {answer.get('substitution')}")
                if answer.get("calculation"):
                    parts.append(f"Calculation:\n{answer.get('calculation')}")
                if answer.get("final_answer"):
                    parts.append(f"{answer.get('final_answer')}")
                return "\n".join(parts)
            
            # Original structured format
            parts = []
            if answer.get("given"):
                parts.append(f"Given: {answer.get('given')}")
            if answer.get("definition"):
                parts.append(f"Definition: {answer.get('definition')}")
            if answer.get("formula"):
                parts.append(f"Formula/Theorem: {answer.get('formula')}")
            if answer.get("coefficients"):
                parts.append(f"Coefficients: {answer.get('coefficients')}")
            if answer.get("steps") and isinstance(answer.get("steps"), list):
                parts.append("Step-by-step Working:")
                for i, step in enumerate(answer.get("steps", []), 1):
                    parts.append(f"Step {i}: {step}")
            if answer.get("function_values") and isinstance(answer.get("function_values"), list):
                parts.append("Function Values:")
                for value in answer.get("function_values", []):
                    parts.append(f"  • {value}")
            if answer.get("final"):
                parts.append(f"Final Conclusion: {answer.get('final')}")
            return "\n".join(parts)
        return str(answer) if answer else ""
    
    # Verify font is registered before using it
    if font_name not in pdfmetrics.getRegisteredFontNames():
        print(f"⚠️  Font {font_name} not registered, attempting to register now...")
        base_dir = os.path.join(os.path.dirname(__file__), "fonts")
        font_file = FONT_NAME_MAP.get(font_name, "NotoSans-Regular.ttf")
        font_path = os.path.join(base_dir, font_file)
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                print(f"✅ Successfully registered font: {font_name}")
            except Exception as e:
                print(f"❌ Failed to register font {font_name}: {e}")
                # Fallback to Helvetica if font registration fails
                font_name = "Helvetica"
        else:
            print(f"❌ Font file not found: {font_path}, using Helvetica")
            font_name = "Helvetica"
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch
    )
    story = []

    def wrap_text(text: str, font: str, bold: bool = False, italic: bool = False, size: int = None) -> str:
        text = normalize_text(str(text))
        escaped = html.escape(text, quote=False)
        bold_tag = "<b>" if bold else ""
        bold_close = "</b>" if bold else ""
        italic_tag = "<i>" if italic else ""
        italic_close = "</i>" if italic else ""
        size_tag = f" size='{size}'" if size else ""
        return f"<font name='{font}'{size_tag}>{bold_tag}{italic_tag}{escaped}{italic_close}{bold_close}</font>"

    styles = getSampleStyleSheet()
    normal_style = ParagraphStyle(
        'BaseNormal', parent=styles['Normal'],
        fontName=font_name, fontSize=15, leading=22.5,
        spaceAfter=6, alignment=text_alignment
    )
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Heading1'],
        fontName=font_name, fontSize=24, leading=30,
        textColor=colors.HexColor('#003366'),
        spaceAfter=36, alignment=TA_CENTER
    )
    question_style = ParagraphStyle(
        'Question', parent=normal_style,
        fontName=font_name, fontSize=15, leading=22.5,
        spaceAfter=12, leftIndent=12 if not is_rtl_lang else 0,
        rightIndent=0 if not is_rtl_lang else 12, alignment=text_alignment
    )
    answer_style = ParagraphStyle(
        'Answer', parent=normal_style,
        fontName=font_name, fontSize=15, leading=22.5,
        textColor=colors.HexColor('#047857'), spaceAfter=12,
        leftIndent=15 if not is_rtl_lang else 0,
        rightIndent=0 if not is_rtl_lang else 15,
        backColor=colors.HexColor('#d1fae5'), alignment=text_alignment
    )

    # Verify font is available before building PDF
    available_fonts = pdfmetrics.getRegisteredFontNames()
    if font_name not in available_fonts:
        print(f"⚠️  WARNING: Font '{font_name}' not in registered fonts: {available_fonts}")
        print(f"⚠️  This will cause black blocks! Attempting to register now...")
        # Try to register the font again
        base_dir = os.path.join(os.path.dirname(__file__), "fonts")
        font_file = FONT_NAME_MAP.get(font_name, "NotoSans-Regular.ttf")
        font_path = os.path.join(base_dir, font_file)
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                print(f"✅ Successfully registered font: {font_name}")
            except Exception as e:
                print(f"❌ Failed to register font: {e}")
                # Use Helvetica as last resort
                font_name = "Helvetica"
        else:
            print(f"❌ Font file not found: {font_path}")
            font_name = "Helvetica"
    
    title = normalize_text(title)
    story.append(Paragraph(wrap_text(title, font_name, bold=True, size=24), title_style))
    story.append(Spacer(1, 0.4 * inch))

    questions = qna_data.get("questions", [])
    for idx, q in enumerate(questions, 1):
        question_text_raw = normalize_text(q.get("question", ""))
        q_text = f"{wrap_text(f'Q{idx}.', font_name, bold=True, size=15)} {wrap_text(question_text_raw, font_name, bold=True, size=15)}"

        if q.get("type") == "mcq" and q.get("options"):
            q_text += "<br/><br/>"
            for i, opt in enumerate(q.get("options", [])):
                opt = normalize_text(opt)
                opt_label = wrap_text(chr(65 + i) + ".", font_name, bold=True, size=14)
                opt_html = wrap_text(opt, font_name, size=14)
                q_text += f"{opt_label} {opt_html}<br/>"

        if q.get("marks"):
            marks_raw = normalize_text(str(q.get("marks")))
            q_text += f"<br/>{wrap_text(f'Marks: {marks_raw}', font_name, italic=True, size=13)}"

        question_table = Table([[Paragraph(q_text, question_style)]], colWidths=[6.8 * inch])
        table_alignment = 'RIGHT' if is_rtl_lang else 'LEFT'
        question_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.white),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), table_alignment),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 15),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 18),
            ('TOPPADDING', (0, 0), (-1, -1), 18),
            ('LEFTPADDING', (0, 0), (-1, -1), 18),
            ('RIGHTPADDING', (0, 0), (-1, -1), 18),
            ('GRID', (0, 0), (-1, -1), 1.5, colors.HexColor('#9ca3af')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(question_table)
        story.append(Spacer(1, 0.25 * inch))

        if output_format in ["questions_answers", "answers_only"]:
            # Try correct_answer first, then answer field as fallback
            correct_answer_raw = q.get("correct_answer") or q.get("answer")
            
            # Only add answer if it exists and is not empty
            if correct_answer_raw and correct_answer_raw != "N/A" and str(correct_answer_raw).strip():
                answer_label = "✓ Answer:" if not is_rtl_lang else ":الإجابة ✓"
                answer_label = normalize_text(answer_label)
                
                # Use format_structured_answer function for consistent formatting
                correct_raw = normalize_text(format_structured_answer(correct_answer_raw, q.get("marks", 0)))
                
                # Only add if formatted answer is not empty
                if correct_raw and correct_raw.strip():
                    answer_text = f"{wrap_text(answer_label, font_name, bold=True, size=15)} {wrap_text(correct_raw, font_name, bold=True, size=15)}"
                    answer_table = Table([[Paragraph(answer_text, answer_style)]], colWidths=[6.8 * inch])
                    answer_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#d1fae5')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#047857')),
                        ('ALIGN', (0, 0), (-1, -1), table_alignment),
                        ('FONTNAME', (0, 0), (-1, -1), font_name),
                        ('FONTSIZE', (0, 0), (-1, -1), 15),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
                        ('TOPPADDING', (0, 0), (-1, -1), 15),
                        ('LEFTPADDING', (0, 0), (-1, -1), 18),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 18),
                        ('GRID', (0, 0), (-1, -1), 2, colors.HexColor('#10b981')),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ]))
                    story.append(answer_table)
        story.append(Spacer(1, 0.3 * inch))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()

def generate_pdf(
    qna_data: Dict[str, Any], 
    output_format: str, 
    title: str = "Generated Questions",
    target_language: Optional[str] = None
) -> bytes:
    """
    Generate high-quality PDF from Q/A data with multilingual support using pdfkit (wkhtmltopdf).
    Uses HTML/CSS for exact formatting and proper Unicode rendering of Indic scripts and Arabic.
    """

    def normalize_text(text: str) -> str:
        """Normalize Unicode text so Tamil, Hindi, Telugu render correctly."""
        if not isinstance(text, str):
            return text
        return unicodedata.normalize("NFC", text)
    
    def format_structured_answer(answer, marks=0):
        """Convert structured answer (dict) to formatted text string"""
        if isinstance(answer, dict):
            # Check if it's board-style format (post-processed 10-mark math)
            if answer.get("substitution") is not None or answer.get("calculation") is not None or answer.get("final_answer") is not None:
                # Board-style format: Given, Formula, Substitution, Calculation, Final Answer
                parts = []
                if answer.get("given"):
                    parts.append(f"Given: {answer.get('given')}")
                if answer.get("formula"):
                    parts.append(f"Formula: {answer.get('formula')}")
                if answer.get("substitution"):
                    parts.append(f"Substitution: {answer.get('substitution')}")
                if answer.get("calculation"):
                    parts.append(f"Calculation:\n{answer.get('calculation')}")
                if answer.get("final_answer"):
                    parts.append(f"{answer.get('final_answer')}")
                return "\n".join(parts)
            
            # Original structured format
            parts = []
            if answer.get("given"):
                parts.append(f"Given: {answer.get('given')}")
            if answer.get("definition"):
                parts.append(f"Definition: {answer.get('definition')}")
            if answer.get("formula"):
                parts.append(f"Formula/Theorem: {answer.get('formula')}")
            if answer.get("coefficients"):
                parts.append(f"Coefficients: {answer.get('coefficients')}")
            if answer.get("steps") and isinstance(answer.get("steps"), list):
                parts.append("Step-by-step Working:")
                for i, step in enumerate(answer.get("steps", []), 1):
                    parts.append(f"Step {i}: {step}")
            if answer.get("function_values") and isinstance(answer.get("function_values"), list):
                parts.append("Function Values:")
                for value in answer.get("function_values", []):
                    parts.append(f"  • {value}")
            if answer.get("final"):
                parts.append(f"Final Conclusion: {answer.get('final')}")
            return "\n".join(parts)
        return str(answer) if answer else ""

    # Detect language
    if not target_language:
        sample_text = ""
        questions = qna_data.get("questions", [])
        if questions:
            first_q = questions[0]
            sample_text = normalize_text(
                f"{first_q.get('question', '')} {format_structured_answer(first_q.get('correct_answer', ''), 0)}"
            )
            if first_q.get("options"):
                sample_text += " " + " ".join(
                    normalize_text(opt) for opt in first_q.get("options", [])
                )
        target_language = detect_language(sample_text, "english")

    font_name = LANGUAGE_TO_FONT_NAME.get(target_language, "NotoLatin")
    is_rtl_lang = is_rtl(target_language)
    
    # Get font file path
    base_dir = os.path.join(os.path.dirname(__file__), "fonts")
    font_file = FONT_NAME_MAP.get(font_name, "NotoSans-Regular.ttf")
    font_path = os.path.abspath(os.path.join(base_dir, font_file))
    
    # Normalize title
    title = normalize_text(title)
    
    # Build HTML content
    html_parts = []
    
    # CSS with @font-face for proper font embedding
    # Prepare font URLs for different rendering engines
    font_url = font_path.replace('\\', '/')
    if os.name == 'nt':  # Windows
        font_url_file = f"file:///{font_url}"
    else:
        font_url_file = f"file://{font_url}"
    
    # Also prepare relative path
    font_url_relative = f"fonts/{font_file}"
    
    # Try to load font as base64 for better compatibility (Playwright)
    font_base64 = None
    if os.path.exists(font_path):
        try:
            import base64
            with open(font_path, 'rb') as f:
                font_data = f.read()
                font_base64 = base64.b64encode(font_data).decode('utf-8')
            print(f"✅ Font loaded as base64: {font_name} ({len(font_base64)} chars)")
        except Exception as e:
            print(f"⚠️  Failed to load font as base64: {e}")
            pass  # Fallback to file URL
    else:
        print(f"❌ Font file not found: {font_path}")
    
    # Build font src with fallbacks - prefer base64 for maximum compatibility
    if font_base64:
        # Base64 encoding works best for all PDF generators
        font_src = f"url(data:font/truetype;charset=utf-8;base64,{font_base64}) format('truetype')"
    else:
        # Fallback to file URLs
        font_src = f"url('{font_url_file}') format('truetype'), url('{font_url_relative}') format('truetype')"
    
    css_content = f"""
    @page {{
        size: A4;
        margin: 1in;
    }}
    
    @font-face {{
        font-family: '{font_name}';
        src: {font_src};
        font-weight: normal;
        font-style: normal;
        font-display: swap;
    }}
    
    /* Force font application */
    body, * {{
        font-family: '{font_name}' !important;
        box-sizing: border-box;
        -webkit-font-smoothing: antialiased;
        -moz-osx-font-smoothing: grayscale;
    }}
    
    body {{
        font-size: 12pt;
        line-height: 1.8;
        color: #000000;
        direction: {'rtl' if is_rtl_lang else 'ltr'};
        text-align: {'right' if is_rtl_lang else 'left'};
        margin: 0;
        padding: 0;
        background: #ffffff;
    }}
    
    .title {{
        font-size: 18pt;
        font-weight: 700;
        color: #000000;
        text-align: center;
        margin-bottom: 30pt;
        margin-top: 0;
        text-transform: uppercase;
        letter-spacing: 1pt;
        line-height: 1.3;
    }}
    
    .question-box {{
        border: none;
        padding: 0;
        margin-bottom: 20pt;
        background: transparent;
        box-shadow: none;
    }}
    
    .question-label {{
        font-weight: 700;
        font-size: 12pt;
        margin-right: 8pt;
        color: #000000;
        display: inline;
    }}
    
    .question-text {{
        font-weight: 400;
        font-size: 12pt;
        margin-bottom: 12pt;
        line-height: 1.8;
        color: #000000;
        display: inline;
    }}
    
    .options {{
        margin: 8pt 0 8pt 20pt;
        padding: 0;
        list-style: none;
    }}
    
    .option {{
        margin: 6pt 0;
        font-size: 12pt;
        line-height: 1.8;
        color: #000000;
        padding: 0;
    }}
    
    .option-label {{
        font-weight: 400;
        color: #000000;
        margin-right: 8pt;
    }}
    
    .marks {{
        font-style: normal;
        font-size: 10pt;
        color: #000000;
        margin-top: 8pt;
        margin-left: 20pt;
        font-weight: 400;
    }}
    
    .answer-box {{
        border: none;
        padding: 0;
        margin-bottom: 20pt;
        margin-top: 10pt;
        background: transparent;
        box-shadow: none;
    }}
    
    .answer-label {{
        font-weight: 400;
        font-size: 12pt;
        color: #000000;
        margin-right: 8pt;
        display: inline;
    }}
    
    .answer-text {{
        font-weight: 400;
        font-size: 12pt;
        color: #000000;
        line-height: 1.8;
        margin: 0;
        display: inline;
    }}
    """
    
    # HTML structure
    html_parts.append(f"""<!DOCTYPE html>
<html lang="{target_language}" dir="{'rtl' if is_rtl_lang else 'ltr'}">
<head>
    <meta charset="UTF-8">
    <style>{css_content}</style>
</head>
<body>
    <div class="title">{html.escape(title)}</div>
""")
    
    questions = qna_data.get("questions", [])
    
    for idx, q in enumerate(questions, 1):
        question_text = normalize_text(q.get("question", ""))
        
        html_parts.append('<div class="question-box">')
        html_parts.append(f'<span class="question-label">Q{idx}.</span>')
        html_parts.append(f'<div class="question-text">{html.escape(question_text)}</div>')
        
        # MCQ Options
        if q.get("type") == "mcq" and q.get("options"):
            html_parts.append('<div class="options">')
            for i, opt in enumerate(q.get("options", [])):
                opt = normalize_text(opt)
                option_label = chr(65 + i)  # A, B, C, D
                html_parts.append(
                    f'<div class="option">'
                    f'<span class="option-label">{option_label}.</span>'
                    f'<span>{html.escape(opt)}</span>'
                    f'</div>'
                )
            html_parts.append('</div>')
        
        # Marks
        if q.get("marks"):
            marks = normalize_text(str(q.get("marks")))
            html_parts.append(f'<div class="marks">Marks: {html.escape(marks)}</div>')
        
        html_parts.append('</div>')
        
        # Answer
        if output_format in ["questions_answers", "answers_only"]:
            answer_label = "✓ Answer:" if not is_rtl_lang else ":الإجابة ✓"
            answer_label = normalize_text(answer_label)
            # Try correct_answer first, then answer field as fallback
            correct_answer_raw = q.get("correct_answer") or q.get("answer") or "N/A"
            
            # Only add answer if it's not "N/A" and not empty
            if correct_answer_raw and correct_answer_raw != "N/A" and str(correct_answer_raw).strip():
                correct_answer = normalize_text(format_structured_answer(correct_answer_raw, q.get("marks", 0)))
                
                # Only add if formatted answer is not empty
                if correct_answer and correct_answer.strip():
                    html_parts.append('<div class="answer-box">')
                    html_parts.append(f'<div class="answer-label">{html.escape(answer_label)}</div>')
                    html_parts.append(f'<div class="answer-text">{html.escape(correct_answer)}</div>')
                    html_parts.append('</div>')
    
    html_parts.append('</body></html>')
    
    html_string = '\n'.join(html_parts)
    
    # Note: Playwright async is handled in the router endpoint
    # Skip Playwright here since we're in sync context
    
    # Try xhtml2pdf (pure Python, good rendering)
    if XHTML2PDF_AVAILABLE:
        try:
            print("📄 Using xhtml2pdf for PDF generation...")
            pdf_bytes = _generate_pdf_xhtml2pdf(html_string, font_path, font_name)
            print("✅ xhtml2pdf PDF generated successfully!")
            return pdf_bytes
        except Exception as e:
            print(f"⚠️  xhtml2pdf error: {e}, trying pdfkit...")
            import traceback
            traceback.print_exc()
    
    # Try pdfkit (requires wkhtmltopdf binary)
    if not PDFKIT_AVAILABLE:
        print("⚠️  pdfkit not installed, using ReportLab fallback...")
        print("⚠️  NOTE: ReportLab may have font rendering issues. Install Playwright for best quality!")
        return _generate_pdf_reportlab(qna_data, output_format, title, target_language, normalize_text, font_name, is_rtl_lang, text_alignment)
    
    try:
        # pdfkit options for better quality and rendering
        # IMPORTANT: enable-local-file-access must be empty string (not None) for fonts to load
        options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': "",  # CRITICAL: Empty string (not None) to allow local font files
            'disable-smart-shrinking': None,  # Prevent text shrinking
            'print-media-type': None,  # Use print media CSS
            'dpi': 300,  # Higher DPI for better quality
            'image-dpi': 300,  # Higher image DPI
            'image-quality': 94,  # High image quality
        }
        
        # Try to find wkhtmltopdf binary
        config = None
        if os.name == 'nt':  # Windows
            possible_paths = [
                r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe',
                r'C:\Program Files (x86)\wkhtmltopdf\bin\wkhtmltopdf.exe',
            ]
            for path in possible_paths:
                if os.path.exists(path):
                    config = pdfkit.configuration(wkhtmltopdf=path)
                    break
        
        # Generate PDF with pdfkit
        if config:
            pdf_bytes = pdfkit.from_string(html_string, False, options=options, configuration=config)
        else:
            pdf_bytes = pdfkit.from_string(html_string, False, options=options)
        
        return pdf_bytes
        
    except (OSError, ImportError, Exception) as e:
        # Fallback to ReportLab if pdfkit/wkhtmltopdf not available
        print(f"⚠️  pdfkit not available ({e}), falling back to ReportLab...")
        return _generate_pdf_reportlab(qna_data, output_format, title, target_language, normalize_text, font_name, is_rtl_lang, text_alignment)

def generate_docx(qna_data: Dict[str, Any], output_format: str, title: str = "Generated Questions") -> bytes:
    """
    Generate DOCX from Q/A data
    """
    # Local safety fallback in case module-level helper is unavailable in runtime
    def _fmt(answer, marks=0):
        try:
            return format_structured_answer(answer, marks)
        except NameError:
            if isinstance(answer, dict):
                if (
                    answer.get("substitution") is not None
                    or answer.get("calculation") is not None
                    or answer.get("final_answer") is not None
                ):
                    parts = []
                    if answer.get("given"):
                        parts.append(f"Given: {answer.get('given')}")
                    if answer.get("formula"):
                        parts.append(f"Formula: {answer.get('formula')}")
                    if answer.get("substitution"):
                        parts.append(f"Substitution: {answer.get('substitution')}")
                    if answer.get("calculation"):
                        parts.append(f"Calculation:\n{answer.get('calculation')}")
                    if answer.get("final_answer"):
                        parts.append(f"{answer.get('final_answer')}")
                    return "\n".join(parts)
                parts = []
                if answer.get("given"):
                    parts.append(f"Given: {answer.get('given')}")
                if answer.get("definition"):
                    parts.append(f"Definition: {answer.get('definition')}")
                if answer.get("formula"):
                    parts.append(f"Formula/Theorem: {answer.get('formula')}")
                if answer.get("coefficients"):
                    parts.append(f"Coefficients: {answer.get('coefficients')}")
                if answer.get("steps") and isinstance(answer.get("steps"), list):
                    parts.append("Step-by-step Working:")
                    for i, step in enumerate(answer.get("steps", []), 1):
                        parts.append(f"Step {i}: {step}")
                if answer.get("function_values") and isinstance(answer.get("function_values"), list):
                    parts.append("Function Values:")
                    for value in answer.get("function_values", []):
                        parts.append(f"  • {value}")
                if answer.get("final"):
                    parts.append(f"Final Conclusion: {answer.get('final')}")
                return "\n".join(parts)
            return str(answer) if answer else ""

    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = 1  # Center
    
    questions = qna_data.get("questions", [])
    
    for idx, q in enumerate(questions, 1):
        # Question with formatting
        q_para = doc.add_paragraph()
        q_para.add_run(f"Q{idx}. ").bold = True
        q_para.add_run(q.get('question', ''))
        q_para_format = q_para.paragraph_format
        q_para_format.left_indent = Inches(0.2)
        q_para_format.space_after = Pt(6)
        
        # MCQ Options with A, B, C, D labels
        if q.get('type') == 'mcq' and q.get('options'):
            for i, opt in enumerate(q.get('options', [])):
                option_label = chr(65 + i)  # A, B, C, D
                opt_para = doc.add_paragraph()
                opt_para.add_run(f"{option_label}. ").bold = True
                opt_para.add_run(opt)
                opt_para_format = opt_para.paragraph_format
                opt_para_format.left_indent = Inches(0.5)
                opt_para_format.space_after = Pt(3)
        
        # Marks display
        if q.get('marks'):
            marks_para = doc.add_paragraph()
            marks_para.add_run(f"Marks: {q.get('marks')} mark{'s' if q.get('marks') != 1 else ''}").italic = True
            marks_para_format = marks_para.paragraph_format
            marks_para_format.left_indent = Inches(0.5)
            marks_para_format.space_after = Pt(6)
        
        # Answer (if format includes answers)
        if output_format in ["questions_answers", "answers_only"]:
            # Try correct_answer first, then answer field as fallback
            correct_answer_raw = q.get('correct_answer') or q.get('answer')
            
            # Only add answer if it exists and is not empty
            if correct_answer_raw and correct_answer_raw != 'N/A' and str(correct_answer_raw).strip():
                answer_text = _fmt(correct_answer_raw, q.get('marks', 0))
                
                # Only add if formatted answer is not empty
                if answer_text and answer_text.strip():
                    answer_para = doc.add_paragraph()
                    answer_para.add_run("✓ Answer: ").bold = True
                    answer_para.add_run(answer_text)
            answer_para_format = answer_para.paragraph_format
            answer_para_format.left_indent = Inches(0.5)
            answer_para_format.space_before = Pt(6)
            answer_para_format.space_after = Pt(12)
        
        doc.add_paragraph()  # Spacing between questions
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def generate_txt(qna_data: Dict[str, Any], output_format: str, title: str = "Generated Questions") -> str:
    """
    Generate TXT from Q/A data
    """
    def _fmt(answer, marks=0):
        try:
            return format_structured_answer(answer, marks)
        except NameError:
            if isinstance(answer, dict):
                if (
                    answer.get("substitution") is not None
                    or answer.get("calculation") is not None
                    or answer.get("final_answer") is not None
                ):
                    parts = []
                    if answer.get("given"):
                        parts.append(f"Given: {answer.get('given')}")
                    if answer.get("formula"):
                        parts.append(f"Formula: {answer.get('formula')}")
                    if answer.get("substitution"):
                        parts.append(f"Substitution: {answer.get('substitution')}")
                    if answer.get("calculation"):
                        parts.append(f"Calculation:\n{answer.get('calculation')}")
                    if answer.get("final_answer"):
                        parts.append(f"{answer.get('final_answer')}")
                    return "\n".join(parts)
                parts = []
                if answer.get("given"):
                    parts.append(f"Given: {answer.get('given')}")
                if answer.get("definition"):
                    parts.append(f"Definition: {answer.get('definition')}")
                if answer.get("formula"):
                    parts.append(f"Formula/Theorem: {answer.get('formula')}")
                if answer.get("coefficients"):
                    parts.append(f"Coefficients: {answer.get('coefficients')}")
                if answer.get("steps") and isinstance(answer.get("steps"), list):
                    parts.append("Step-by-step Working:")
                    for i, step in enumerate(answer.get("steps", []), 1):
                        parts.append(f"Step {i}: {step}")
                if answer.get("function_values") and isinstance(answer.get("function_values"), list):
                    parts.append("Function Values:")
                    for value in answer.get("function_values", []):
                        parts.append(f"  • {value}")
                if answer.get("final"):
                    parts.append(f"Final Conclusion: {answer.get('final')}")
                return "\n".join(parts)
            return str(answer) if answer else ""

    lines = []
    lines.append("=" * 60)
    lines.append(title.center(60))
    lines.append("=" * 60)
    lines.append("")
    
    questions = qna_data.get("questions", [])
    
    for idx, q in enumerate(questions, 1):
        lines.append(f"Q{idx}. {q.get('question', '')}")
        lines.append("")
        
        if q.get('type') == 'mcq' and q.get('options'):
            for opt in q.get('options', []):
                lines.append(f"    {opt}")
            lines.append("")
        
        if output_format in ["questions_answers", "answers_only"]:
            # Try correct_answer first, then answer field as fallback
            correct_answer_raw = q.get('correct_answer') or q.get('answer')
            
            # Only add answer if it exists and is not empty
            if correct_answer_raw and correct_answer_raw != 'N/A' and str(correct_answer_raw).strip():
                answer_text = _fmt(correct_answer_raw, q.get('marks', 0))
                
                # Only add if formatted answer is not empty
                if answer_text and answer_text.strip():
                    lines.append(f"Answer: {answer_text}")
                    lines.append("")
        
        lines.append("-" * 60)
        lines.append("")
    
    return "\n".join(lines)

