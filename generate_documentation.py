"""
Generate comprehensive documentation for StudyQnA Generator application
Creates a single .doc file with all documentation
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def create_documentation():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('StudyQnA Generator', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Complete Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.paragraph_format
    subtitle_format.space_after = Pt(12)
    
    date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Introduction',
        '2. Features Overview',
        '3. Advantages',
        '4. Software & Technologies Used',
        '5. User Dashboard - Features & Options',
        '6. Admin Dashboard - Features & Options',
        '7. Security Measures',
        '8. Security Constraints',
        '9. User Journey',
        '10. Application Architecture',
        '11. API Documentation',
        '12. Installation Guide',
        '13. Deployment Guide',
        '14. Limitations',
        '15. Future Enhancements',
        '16. Screenshots Structure'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.25)
    
    doc.add_page_break()
    
    # 1. Introduction
    doc.add_heading('1. Introduction', 1)
    doc.add_paragraph(
        'StudyQnA Generator is a secure, production-ready web application designed to help students and '
        'teachers create high-quality question and answer sets from educational materials. The application '
        'supports multiple input formats (PDF and images), uses AI to generate questions, and provides '
        'multilingual support for various languages including Tamil, Hindi, Telugu, Kannada, Malayalam, '
        'Arabic, Spanish, and English.'
    )
    doc.add_paragraph(
        'The system implements a robust security framework with role-based access control, content validation, '
        'and comprehensive audit logging. It supports both free and premium user tiers with different quota limits, '
        'and includes an admin dashboard for managing users, monitoring usage, and maintaining system security.'
    )
    
    # 2. Features Overview
    doc.add_heading('2. Features Overview', 1)
    
    doc.add_heading('2.1 Core Features', 2)
    features = [
        'OTP-based Email Authentication: Secure login system using one-time passwords sent via email',
        'Multi-format File Upload: Support for PDF documents and image files (JPG, PNG)',
        'OCR Text Extraction: Automatic text extraction from uploaded PDFs and images using Tesseract OCR',
        'AI-Powered Q/A Generation: Generate questions and answers using OpenAI GPT models',
        'Multiple Question Types: Support for Multiple Choice Questions (MCQ) and Descriptive questions',
        'Mark Pattern Selection: Choose from predefined mark patterns (1, 2, 3, 5, 10, Mixed)',
        'Difficulty Levels: Select question difficulty (Easy, Medium, Hard)',
        'Multilingual Support: Generate questions in 8 languages (Tamil, Hindi, Telugu, Kannada, Malayalam, Arabic, Spanish, English)',
        'Saved Q/A Sets: Save and manage generated question sets',
        'Multiple Download Formats: Export as PDF, DOCX, or TXT files',
        'Premium Welcome Animation: Confetti animation when premium is activated',
        'User Review System: Submit ratings and feedback (1-5 stars with optional text)',
        'Mobile Camera Support: Capture images directly from mobile devices',
        'Content Validation: Strict filtering to ensure only educational content is uploaded',
        'Usage Statistics: Track PDF and image upload quotas',
        'Monthly Quota Reset: Automatic quota reset based on premium activation date'
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('2.2 User Features', 2)
    user_features = [
        'Free Tier: 3 PDFs/month, 5 images/month, max 3 questions per generation',
        'Premium Tier: 15 PDFs/month, 20 images/month, max 20 questions per generation',
        'Profile Management: View user details, premium status, and usage statistics',
        'Upload History: View all uploaded files with preview options',
        'Q/A Set Management: View, download, and manage saved question sets',
        'Teacher Mode: Special mode for educators with enhanced features'
    ]
    
    for feature in user_features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('2.3 Admin Features', 2)
    admin_features = [
        'User Management: View all users, approve/reject premium requests',
        'Usage Monitoring: Track AI token usage with threshold alerts',
        'Upload Management: View all uploads with preview capabilities',
        'Q/A Set Management: View all generated question sets',
        'Review Management: View and delete user reviews',
        'Login Logs: Monitor user login/logout activities with IP addresses and device types',
        'Audit Logging: Track all system activities',
        'Quota Management: Adjust user quotas manually',
        'Export Logs: Export logs in daily, monthly, or yearly formats'
    ]
    
    for feature in admin_features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 3. Advantages
    doc.add_heading('3. Advantages', 1)
    advantages = [
        'Time-Saving: Automatically generates questions from educational materials, saving hours of manual work',
        'Multilingual Support: Supports 8 languages, making it accessible to diverse student populations',
        'High Quality: Uses advanced AI models to generate realistic exam-style questions',
        'Secure: Comprehensive security measures protect user data and content',
        'Scalable: Built with modern technologies that can handle growing user bases',
        'User-Friendly: Intuitive interface for both students and teachers',
        'Mobile-Friendly: Responsive design with mobile camera support',
        'Flexible: Multiple question types, mark patterns, and difficulty levels',
        'Cost-Effective: Free tier available, premium tier offers excellent value',
        'Professional Output: Generates high-quality PDFs matching exam paper standards',
        'Content Safety: Strict validation ensures only appropriate educational content',
        'Audit Trail: Complete logging for accountability and security'
    ]
    
    for advantage in advantages:
        p = doc.add_paragraph(advantage, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 4. Software & Technologies Used
    doc.add_heading('4. Software & Technologies Used', 1)
    
    doc.add_heading('4.1 Frontend Technologies', 2)
    frontend_tech = [
        'React 18+: Modern JavaScript library for building user interfaces',
        'Tailwind CSS: Utility-first CSS framework for rapid UI development',
        'Vite: Fast build tool and development server',
        'Axios: HTTP client for API requests',
        'React Router DOM: Client-side routing',
        'Canvas Confetti: Animation library for premium welcome effects',
        'React Hot Toast: Toast notification library',
        'PDF.js: PDF rendering library for previews'
    ]
    
    for tech in frontend_tech:
        p = doc.add_paragraph(tech, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('4.2 Backend Technologies', 2)
    backend_tech = [
        'FastAPI: Modern Python web framework for building APIs',
        'Python 3.13: Latest Python version',
        'SQLAlchemy: ORM for database operations',
        'PostgreSQL: Robust relational database',
        'Alembic: Database migration tool',
        'JWT: JSON Web Tokens for authentication',
        'Pydantic: Data validation using Python type annotations',
        'OpenAI API: AI-powered question generation',
        'Tesseract OCR: Optical Character Recognition for text extraction',
        'PyPDF2: PDF processing library',
        'Pillow (PIL): Image processing library',
        'OpenCV: Computer vision library for content validation',
        'Ultralytics YOLO: Human detection for content filtering',
        'ReportLab: PDF generation library',
        'Playwright: Browser automation for high-quality PDF generation',
        'python-docx: DOCX file generation',
        'aiosmtplib: Asynchronous SMTP client for emails',
        'email-validator: Email validation library',
        'dnspython: DNS lookup for email validation'
    ]
    
    for tech in backend_tech:
        p = doc.add_paragraph(tech, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('4.3 Database', 2)
    doc.add_paragraph('PostgreSQL: Enterprise-grade open-source relational database management system')
    
    doc.add_heading('4.4 Deployment & Infrastructure', 2)
    deployment_tech = [
        'Docker: Containerization platform',
        'Docker Compose: Multi-container Docker application management',
        'Nginx: Web server and reverse proxy',
        'AWS Lightsail: Cloud hosting platform (for V1)'
    ]
    
    for tech in deployment_tech:
        p = doc.add_paragraph(tech, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 5. User Dashboard - Features & Options
    doc.add_heading('5. User Dashboard - Features & Options', 1)
    
    doc.add_heading('5.1 Navigation Tabs', 2)
    user_tabs = {
        'Dashboard': 'Main landing page showing quota usage, recent uploads, and quick actions',
        'Upload': 'Upload PDFs or images for question generation',
        'Generate': 'Generate questions from uploaded files with customizable options',
        'Saved Sets': 'View and manage all saved question sets',
        'Profile': 'View user information, premium status, and usage statistics',
        'Reviews': 'Submit feedback and ratings for the application'
    }
    
    for tab, desc in user_tabs.items():
        p = doc.add_paragraph(f'{tab}: {desc}', style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('5.2 Upload Options', 2)
    upload_options = [
        'File Type Selection: Choose between PDF or Image upload',
        'File Upload: Drag-and-drop or click to select files',
        'Mobile Camera: Direct camera access on mobile devices (hidden on desktop)',
        'File Validation: Automatic validation of file type, size, and content',
        'Content Filtering: Blocks inappropriate content (human bodies, nudity, violence, etc.)',
        'Preview: Preview uploaded files before processing'
    ]
    
    for option in upload_options:
        p = doc.add_paragraph(option, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('5.3 Generation Options', 2)
    gen_options = [
        'Number of Questions: Type or select number (max 3 for free, max 20 for premium)',
        'Question Type: Choose MCQ or Descriptive',
        'Mark Pattern: Select from 1, 2, 3, 5, 10, or Mixed marks',
        'Difficulty Level: Easy, Medium, or Hard',
        'Target Language: Select from 8 supported languages',
        'Output Format: Questions only, Questions with Answers, or Answers only'
    ]
    
    for option in gen_options:
        p = doc.add_paragraph(option, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('5.4 Profile Page', 2)
    profile_features = [
        'User Email: Display registered email address',
        'Premium Status: Show current premium subscription status',
        'Premium Validity: Display premium expiration date',
        'Monthly Reset Date: Show the day of month when quota resets',
        'Usage Statistics: Visual quota bars for PDFs and Images',
        'Upload Count: Current usage vs. total quota',
        'Account Creation Date: When the account was created'
    ]
    
    for feature in profile_features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 6. Admin Dashboard - Features & Options
    doc.add_heading('6. Admin Dashboard - Features & Options', 1)
    
    doc.add_heading('6.1 Navigation Tabs', 2)
    admin_tabs = {
        'Dashboard': 'Overview of system statistics and recent activities',
        'Users': 'Manage all users, view details, approve/reject premium requests',
        'Uploads': 'View all uploaded files with preview capabilities',
        'Q/A Sets': 'View all generated question sets',
        'Reviews': 'View and manage user reviews',
        'Logs': 'View login logs, audit logs, and system activities',
        'AI Usage': 'Monitor AI token usage and set thresholds',
        'Settings': 'System configuration and settings'
    }
    
    for tab, desc in admin_tabs.items():
        p = doc.add_paragraph(f'{tab}: {desc}', style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('6.2 User Management', 2)
    user_mgmt = [
        'User List: View all registered users with email, status, and premium information',
        'Premium Requests: Approve or reject premium subscription requests',
        'Quota Adjustment: Manually increase or decrease user quotas',
        'User Details: View complete user profile and activity history',
        'Email Display: View user email addresses for communication'
    ]
    
    for feature in user_mgmt:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('6.3 Upload Management', 2)
    upload_mgmt = [
        'Upload List: View all uploaded files across all users',
        'PDF Preview: Preview PDF content using PDF.js',
        'Image Preview: Securely stream and preview images',
        'File Details: View file type, size, upload date, and user information',
        'Content Validation Status: See validation results for each upload'
    ]
    
    for feature in upload_mgmt:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('6.4 Logging & Monitoring', 2)
    logging_features = [
        'Login Logs: Track all user login/logout activities',
        'IP Address Tracking: Record IP addresses for security monitoring',
        'Device Type Detection: Identify if users are on mobile or desktop',
        'Logout Time: Track when users log out',
        'Export Functionality: Export logs in daily, monthly, or yearly formats',
        'AI Usage Tracking: Monitor OpenAI API token consumption',
        'Threshold Alerts: Receive alerts when AI usage reaches configured thresholds',
        'Audit Trail: Complete record of all system activities'
    ]
    
    for feature in logging_features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 7. Security Measures
    doc.add_heading('7. Security Measures', 1)
    
    security_measures = [
        'JWT Authentication: Secure token-based authentication system',
        'OTP-based Login: One-time passwords prevent password-related attacks',
        'Email Validation: RFC-compliant email validation with DNS MX record lookup',
        'Role-Based Access Control (RBAC): Separate user and admin roles with different permissions',
        'Content Validation: Strict filtering of uploaded content to prevent inappropriate material',
        'Human Detection: AI-powered detection to block human body content',
        'File Type Validation: Strict validation of file types and sizes',
        'HTTPS Support: Encrypted communication (in production)',
        'Secure File Storage: Files stored securely on local disk with proper access controls',
        'SQL Injection Prevention: Parameterized queries using SQLAlchemy ORM',
        'XSS Protection: Input sanitization and output encoding',
        'CSRF Protection: Token-based protection against cross-site request forgery',
        'Rate Limiting: Protection against brute force attacks',
        'IP Logging: Track user IP addresses for security monitoring',
        'Audit Logging: Complete audit trail of all system activities',
        'Session Management: Secure session handling with JWT tokens',
        'Password Hashing: Bcrypt hashing for any stored passwords',
        'Input Validation: Comprehensive validation using Pydantic schemas',
        'Error Handling: Secure error messages that don\'t leak system information'
    ]
    
    for measure in security_measures:
        p = doc.add_paragraph(measure, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 8. Security Constraints
    doc.add_heading('8. Security Constraints', 1)
    
    doc.add_heading('8.1 Content Filtering Constraints', 2)
    content_constraints = [
        'BLOCKED: Photos containing full or partial human bodies',
        'BLOCKED: Nudity or inappropriate content',
        'BLOCKED: Violence or graphic content',
        'BLOCKED: Medical body parts',
        'BLOCKED: Personally Identifiable Information (PII)',
        'ALLOWED: Photos of textbook pages',
        'ALLOWED: Printed study materials',
        'ALLOWED: Diagrams, charts, tables, graphs',
        'ALLOWED: Clean text from notebooks or printed sheets',
        'ALLOWED: Educational diagrams and illustrations'
    ]
    
    for constraint in content_constraints:
        p = doc.add_paragraph(constraint, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('8.2 File Upload Constraints', 2)
    file_constraints = [
        'Maximum File Size: 10MB for PDFs, 5MB for images (configurable)',
        'Allowed Formats: PDF, JPG, PNG only',
        'File Type Validation: Strict MIME type checking',
        'Content Scanning: Automatic content validation before acceptance',
        'Quota Limits: Enforced per user tier (free/premium)'
    ]
    
    for constraint in file_constraints:
        p = doc.add_paragraph(constraint, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('8.3 Access Control Constraints', 2)
    access_constraints = [
        'Admin-Only Endpoints: Certain endpoints restricted to admin users only',
        'Premium-Only Features: Download and advanced features require premium subscription',
        'User Isolation: Users can only access their own data',
        'Rate Limiting: API rate limits to prevent abuse',
        'Session Timeout: Automatic session expiration for security'
    ]
    
    for constraint in access_constraints:
        p = doc.add_paragraph(constraint, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('8.4 Data Protection Constraints', 2)
    data_constraints = [
        'No PII Storage: Personally identifiable information is not stored unnecessarily',
        'Encrypted Storage: Sensitive data encrypted at rest',
        'Secure Transmission: All API communications use HTTPS',
        'Data Retention: Configurable data retention policies',
        'Backup Security: Regular secure backups of database'
    ]
    
    for constraint in data_constraints:
        p = doc.add_paragraph(constraint, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 9. User Journey
    doc.add_heading('9. User Journey', 1)
    
    doc.add_heading('9.1 Student User Journey', 2)
    student_journey = [
        '1. Registration/Login: User enters email, receives OTP, verifies and logs in',
        '2. Upload Material: User uploads PDF or image of study material',
        '3. Content Validation: System validates content is appropriate',
        '4. Generate Questions: User selects options (number, type, marks, difficulty, language)',
        '5. Review Questions: User reviews generated questions and answers',
        '6. Save Set: User saves the question set for future reference',
        '7. Download: User downloads questions in preferred format (PDF/DOCX/TXT)',
        '8. Manage Sets: User can view, download, or delete saved sets'
    ]
    
    for step in student_journey:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('9.2 Teacher User Journey', 2)
    teacher_journey = [
        '1. Login: Teacher logs in with email OTP',
        '2. Upload Course Material: Upload multiple PDFs or images',
        '3. Generate Question Bank: Create multiple question sets with different parameters',
        '4. Review & Edit: Review generated questions for accuracy',
        '5. Export: Download question sets in various formats',
        '6. Share: Use exported files for exams or study materials'
    ]
    
    for step in teacher_journey:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('9.3 Admin User Journey', 2)
    admin_journey = [
        '1. Admin Login: Admin logs in with admin credentials',
        '2. Monitor Users: View all users and their activities',
        '3. Manage Premium: Approve or reject premium requests',
        '4. Review Uploads: Monitor uploaded content for compliance',
        '5. Track Usage: Monitor AI usage and system performance',
        '6. Export Logs: Generate reports for analysis',
        '7. Manage Quotas: Adjust user quotas as needed'
    ]
    
    for step in admin_journey:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 10. Application Architecture
    doc.add_heading('10. Application Architecture', 1)
    
    doc.add_heading('10.1 System Architecture', 2)
    doc.add_paragraph(
        'The application follows a three-tier architecture:'
    )
    arch_tiers = [
        'Presentation Layer (Frontend): React-based SPA with Tailwind CSS',
        'Application Layer (Backend): FastAPI RESTful API with Python',
        'Data Layer: PostgreSQL database for persistent storage'
    ]
    
    for tier in arch_tiers:
        p = doc.add_paragraph(tier, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('10.2 Component Structure', 2)
    components = [
        'Frontend Components: Modular React components for each feature',
        'API Routes: RESTful endpoints organized by feature (auth, upload, qna, admin, user)',
        'Database Models: SQLAlchemy ORM models for data persistence',
        'Services: Business logic layer (OCR, AI, email, storage, validation)',
        'Middleware: Authentication, CORS, error handling',
        'Utilities: Helper functions for common operations'
    ]
    
    for component in components:
        p = doc.add_paragraph(component, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('10.3 Data Flow', 2)
    data_flow = [
        '1. User uploads file → Frontend sends to Backend API',
        '2. Backend validates file → Stores in secure location',
        '3. OCR service extracts text → Returns text content',
        '4. AI service generates Q/A → Returns structured data',
        '5. Data saved to database → User can view/download',
        '6. Download service generates PDF/DOCX → Returns file to user'
    ]
    
    for step in data_flow:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 11. API Documentation
    doc.add_heading('11. API Documentation', 1)
    
    doc.add_heading('11.1 Authentication Endpoints', 2)
    auth_endpoints = [
        'POST /api/auth/otp/request: Request OTP for email login',
        'POST /api/auth/otp/verify: Verify OTP and get JWT token',
        'POST /api/auth/logout: Logout user (invalidates token)'
    ]
    
    for endpoint in auth_endpoints:
        p = doc.add_paragraph(endpoint, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('11.2 User Endpoints', 2)
    user_endpoints = [
        'GET /api/user/me: Get current user information',
        'GET /api/user/profile: Get user profile with statistics',
        'GET /api/upload: Get user uploads',
        'POST /api/upload: Upload PDF or image file',
        'GET /api/qna/sets: Get all saved Q/A sets',
        'POST /api/qna/generate: Generate questions from upload',
        'GET /api/qna/sets/{id}/download: Download Q/A set in specified format',
        'POST /api/reviews: Submit user review'
    ]
    
    for endpoint in user_endpoints:
        p = doc.add_paragraph(endpoint, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('11.3 Admin Endpoints', 2)
    admin_endpoints = [
        'GET /api/admin/users: Get all users',
        'GET /api/admin/premium-requests: Get pending premium requests',
        'POST /api/admin/premium-requests/{id}/approve: Approve premium request',
        'POST /api/admin/premium-requests/{id}/reject: Reject premium request',
        'GET /api/admin/uploads: Get all uploads',
        'GET /api/admin/qna-sets: Get all Q/A sets',
        'GET /api/admin/reviews: Get all reviews',
        'DELETE /api/admin/reviews/{id}: Delete review',
        'GET /api/admin/login-logs: Get login logs',
        'GET /api/admin/logs/export: Export logs',
        'GET /api/admin/ai-usage: Get AI usage statistics',
        'POST /api/admin/users/{id}/quota: Adjust user quota'
    ]
    
    for endpoint in admin_endpoints:
        p = doc.add_paragraph(endpoint, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 12. Installation Guide
    doc.add_heading('12. Installation Guide', 1)
    
    doc.add_heading('12.1 Prerequisites', 2)
    prerequisites = [
        'Python 3.13 or higher',
        'Node.js 18+ and npm',
        'PostgreSQL 14+',
        'Visual C++ Build Tools (for Windows, if using Playwright)',
        'Tesseract OCR installed and in PATH'
    ]
    
    for prereq in prerequisites:
        p = doc.add_paragraph(prereq, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('12.2 Backend Installation', 2)
    backend_steps = [
        '1. Navigate to backend directory',
        '2. Create virtual environment: python -m venv venv',
        '3. Activate virtual environment',
        '4. Install dependencies: pip install -r requirements.txt',
        '5. Install Playwright: pip install playwright && playwright install chromium',
        '6. Configure .env file with database credentials and API keys',
        '7. Initialize database: python init_db.py',
        '8. Run migrations: alembic upgrade head',
        '9. Start server: uvicorn app.main:app --reload'
    ]
    
    for step in backend_steps:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('12.3 Frontend Installation', 2)
    frontend_steps = [
        '1. Navigate to frontend directory',
        '2. Install dependencies: npm install',
        '3. Configure .env file with API URL',
        '4. Start development server: npm run dev'
    ]
    
    for step in frontend_steps:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 13. Deployment Guide
    doc.add_heading('13. Deployment Guide', 1)
    
    doc.add_heading('13.1 Docker Deployment', 2)
    docker_steps = [
        '1. Build Docker images: docker-compose build',
        '2. Start services: docker-compose up -d',
        '3. Run migrations: docker-compose exec backend alembic upgrade head',
        '4. Access application at configured domain'
    ]
    
    for step in docker_steps:
        p = doc.add_paragraph(step, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_heading('13.2 Production Considerations', 2)
    production_considerations = [
        'Use HTTPS: Configure SSL certificates for secure communication',
        'Environment Variables: Store sensitive data in environment variables',
        'Database Backups: Set up regular automated backups',
        'Monitoring: Implement logging and monitoring solutions',
        'Load Balancing: Use load balancer for high availability',
        'CDN: Use CDN for static assets',
        'Rate Limiting: Configure rate limits for API endpoints',
        'Security Headers: Set appropriate security headers'
    ]
    
    for consideration in production_considerations:
        p = doc.add_paragraph(consideration, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 14. Limitations
    doc.add_heading('14. Limitations', 1)
    
    limitations = [
        'OCR Accuracy: OCR accuracy depends on image quality and language complexity',
        'AI Dependency: Requires OpenAI API key and internet connection',
        'File Size Limits: Maximum file sizes enforced (10MB PDF, 5MB images)',
        'Quota Restrictions: Free tier has limited quotas',
        'Language Support: Limited to 8 languages currently',
        'Question Quality: AI-generated questions may require manual review',
        'Storage: Local file storage may need expansion for large deployments',
        'Browser Compatibility: Some features require modern browsers',
        'Mobile Features: Camera access requires HTTPS in production'
    ]
    
    for limitation in limitations:
        p = doc.add_paragraph(limitation, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 15. Future Enhancements
    doc.add_heading('15. Future Enhancements', 1)
    
    enhancements = [
        'Additional Languages: Support for more languages',
        'Question Templates: Pre-defined question templates',
        'Bulk Upload: Support for multiple file uploads at once',
        'Question Editing: Allow users to edit generated questions',
        'Collaboration: Share question sets with other users',
        'Analytics Dashboard: Advanced analytics for teachers',
        'Mobile App: Native mobile applications',
        'Offline Mode: Support for offline question generation',
        'Integration: Integration with LMS platforms',
        'Advanced AI: Fine-tuned models for better question quality',
        'Cloud Storage: Integration with cloud storage providers',
        'API Access: Public API for third-party integrations',
        'White-labeling: Customizable branding for institutions',
        'Multi-tenant Support: Support for multiple institutions'
    ]
    
    for enhancement in enhancements:
        p = doc.add_paragraph(enhancement, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # 16. Screenshots Structure
    doc.add_heading('16. Screenshots Structure', 1)
    
    doc.add_paragraph(
        'The following screenshots should be included in the documentation:'
    )
    
    screenshot_sections = [
        'Login Page: OTP request and verification interface',
        'User Dashboard: Main dashboard with quota bars and recent uploads',
        'Upload Page: File upload interface with mobile camera option',
        'Generate Page: Question generation form with all options',
        'Saved Sets: List of saved question sets with download options',
        'Profile Page: User profile with statistics and premium status',
        'Admin Dashboard: Admin overview with system statistics',
        'User Management: User list with premium request management',
        'Upload Management: All uploads with preview capabilities',
        'Logs View: Login logs and audit trail',
        'AI Usage: AI token usage monitoring dashboard',
        'Generated PDF: Sample of generated PDF output'
    ]
    
    for section in screenshot_sections:
        p = doc.add_paragraph(section, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
    
    # Save document
    output_path = 'StudyQnA_Generator_Documentation.docx'
    doc.save(output_path)
    print(f"✅ Documentation generated successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_documentation()


